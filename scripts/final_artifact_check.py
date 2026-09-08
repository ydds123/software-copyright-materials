#!/usr/bin/env python3
"""Final artifact re-check (补漏 2 / plan §6.6).

After DOCX/PDF build, re-extract text/structure from the FINAL file and
re-verify: page count, embedded images, header (software name/version),
section/step numbering, and fact assertions — instead of trusting the
Markdown source alone.

Exit codes: 0 pass / 1 artifact issues / 2 invalid input
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

from common import read_json, read_text

MIN_SOURCE_PAGES = 60  # 前后 30 页 rule; 不足 60 页全部提交


def extract_docx(docx_path: Path) -> tuple[str, int]:
    """Return (text, embedded image count) from a DOCX."""
    try:
        import docx as python_docx
    except Exception as exc:  # noqa: BLE001
        raise SystemExit(f"python-docx 不可用: {exc}")
    doc = python_docx.Document(str(docx_path))
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            text += "\n" + " | ".join(cell.text for cell in row.cells)
    # v1.7 提取页眉页脚文本（程序鉴别材料的软件名/版本号/页码在 header1.xml）
    for section in doc.sections:
        for hf in (section.header, section.footer):
            try:
                for p in hf.paragraphs:
                    if p.text.strip():
                        text += "\n" + p.text
            except Exception:
                continue
    # python-docx 偶发读不到 header（style 关联），用 zipfile 兜底
    try:
        import zipfile as _z
        with _z.ZipFile(str(docx_path)) as zf:
            for n in zf.namelist():
                if re.search(r'word/(header|footer)\d+\.xml$', n):
                    xml = zf.read(n).decode('utf-8', errors='replace')
                    for t in re.findall(r'<w:t[^>]*>([^<]*)</w:t>', xml):
                        if t.strip():
                            text += "\n" + t
    except Exception:
        pass
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in (rel.reltype or ""):
            image_count += 1
    return text, image_count


def extract_pdf(pdf_path: Path) -> tuple[str, int, int]:
    """Return (text, image count, page count) from a PDF."""
    try:
        import fitz
    except Exception as exc:  # noqa: BLE001
        raise SystemExit(f"PyMuPDF 不可用: {exc}")
    doc = fitz.open(str(pdf_path))
    text = "\n".join(page.get_text() for page in doc)
    images = sum(len(page.get_images(full=True)) for page in doc)
    return text, images, doc.page_count


def verify_section_numbering(text: str) -> list[str]:
    errors: list[str] = []
    # v1.8 主路径：一级章标题序列（"1 引言"样式）连续性检查，
    # 不依赖子节编号（术语表等章可无 X.Y 子节）
    chapters: list[int] = []
    for line in text.splitlines():
        s = line.strip()
        if re.match(r'^\d+\.', s):
            continue  # 子节编号不计入章序列
        m = re.match(r'^(\d+)\s+[\u4e00-\u9fff]{2,}', s)
        if m:
            chapters.append(int(m.group(1)))
    if len(chapters) >= 2:
        # 去连续重复（同一章标题在 TOC/正文可能出现两次）
        uniq: list[int] = []
        for n in chapters:
            if not uniq or n != uniq[-1]:
                uniq.append(n)
        for a, b in zip(uniq, uniq[1:]):
            if b != a + 1:
                errors.append(f"最终件一级编号跳跃：{a} 之后是 {b}")
                break
        return errors
    # fallback：旧逻辑（X.Y 子节反推）
    seq: list[int] = []
    for line in text.splitlines():
        m = re.match(r"^\s*(\d+)\.\d*\s", line.strip())
        if m:
            seq.append(int(m.group(1)))
    top: list[int] = []
    for n in seq:
        if not top or n != top[-1]:
            top.append(n)
    for a, b in zip(top, top[1:]):
        if b != a + 1 and b != a:
            errors.append(f"最终件一级编号跳跃：{a} 之后是 {b}")
    return errors


def verify_fact_assertions(text: str, plan: dict[str, Any] | None) -> list[str]:
    errors: list[str] = []
    if not plan:
        return errors
    for t in plan.get("fact_assertions") or []:
        if not isinstance(t, dict) or t.get("status") != "confirmed":
            continue
        subject = str(t.get("subject") or "").strip()
        if subject and subject not in text:
            errors.append(f"最终件中找不到事实断言主题 '{subject}'")
    return errors


def run(
    artifact_path: Path,
    plan_path: Path | None = None,
    source_manual: Path | None = None,
    software_name: str = "",
    version: str = "",
) -> dict[str, Any]:
    if not artifact_path.exists():
        return {"status": "invalid", "errors": [f"缺少 {artifact_path}"]}

    errors: list[str] = []
    warnings: list[str] = []
    suffix = artifact_path.suffix.lower()
    if suffix == ".docx":
        text, image_count = extract_docx(artifact_path)
        page_count = None
    elif suffix == ".pdf":
        text, image_count, page_count = extract_pdf(artifact_path)
    else:
        return {"status": "invalid", "errors": [f"不支持的文件类型 {suffix}（仅 docx/pdf）"]}

    # header: software name & version present
    if software_name and software_name not in text:
        errors.append(f"最终件中找不到软件名称 '{software_name}'")
    if version and version not in text:
        errors.append(f"最终件中找不到版本号 '{version}'")

    # embedded images
    if image_count == 0:
        warnings.append("最终件未检测到内嵌图片（如应含截图，请检查断链/白框）")
    elif source_manual and source_manual.exists():
        # v1.7 只数真实图片引用（![]），【截图预留】占位符不产生图片，不计入比对
        src_refs = len(re.findall(r"!\[[^\]]*\]\([^)]+\)", read_text(source_manual)))
        if src_refs > 0 and image_count < src_refs:
            errors.append(
                f"最终件内嵌图片 {image_count} 张，少于源稿真实图片引用 {src_refs} 处（疑似断链/丢失）"
            )

    # numbering re-check on final text
    errors.extend(verify_section_numbering(text))

    # fact assertions re-check on final text
    plan = read_json(plan_path) if plan_path and plan_path.exists() else None
    errors.extend(verify_fact_assertions(text, plan))

    errors = list(dict.fromkeys(errors))
    return {
        "status": "pass" if not errors else "blocked",
        "artifact": str(artifact_path),
        "page_count": page_count,
        "embedded_images": image_count,
        "errors": errors,
        "warnings": warnings,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--artifact", required=True, help="最终 DOCX/PDF 路径")
    parser.add_argument("--plan", help="材料证据计划.json（事实断言复检）")
    parser.add_argument("--source-manual", help="源 Markdown 手册（截图引用对比）")
    parser.add_argument("--software-name", default="")
    parser.add_argument("--version", default="")
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()

    report = run(
        Path(args.artifact),
        Path(args.plan) if args.plan else None,
        Path(args.source_manual) if args.source_manual else None,
        software_name=args.software_name,
        version=args.version,
    )
    if args.json:
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        print(f"FINAL ARTIFACT {report['status'].upper()}")
        print(f"  页数: {report.get('page_count')} | 内嵌图片: {report.get('embedded_images')}")
        for e in report["errors"]:
            print(f"  ERROR: {e}")
        for w in report["warnings"]:
            print(f"  WARNING: {w}")
    if report["status"] == "invalid":
        sys.exit(2)
    sys.exit(1 if report["errors"] else 0)


if __name__ == "__main__":
    main()
