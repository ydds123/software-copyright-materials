#!/usr/bin/env python3
"""Batch re-check / risk report (方案 v2 阶段 2): 整批复检，输出风险报告。

按决策③：只出风险报告，不自动返工；旧批次材料冻结不动。
按决策④：确定性检查（文件缺失/版本漂移/编号错误/重复粘贴）列硬错误；
          相似度列风险等级（高/中/低），高风险需人工复核。

用法：
  python batch_risk_report.py --workspace "<年份>年软件著作权申请资料目录" [--out <报告目录>] [--json]

输出：
  - 整批复检报告.md（人工阅读）
  - 整批复检报告.json（机器可读）
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

from batch_structure_check import (
    heading_skeleton,
    heading_numbering_issues,
    paragraph_data,
    table_headers,
    normalize_header,
    run as batch_run,
)

DRAFT_MANUAL = Path("草稿") / "操作手册.md"
TASK_FILE = "任务登记.json"

# 正式件候选排除：临时文件、模板迭代稿、副本
EXCLUDE_PATTERNS = ("~$", "新版模板式", "(1)", "copy", "副本")


def extract_docx_text(path: Path) -> str:
    """提取 DOCX 正文（段落 + 表格），近似还原 Markdown 结构。"""
    try:
        import docx  # type: ignore
    except ImportError:
        return ""
    try:
        d = docx.Document(str(path))
    except Exception:
        return ""
    parts: list[str] = []
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            style = (p.style.name or "") if p.style else ""
            prefix = ""
            if "Heading 1" in style:
                prefix = "# "
            elif "Heading 2" in style:
                prefix = "## "
            elif "Heading 3" in style:
                prefix = "### "
            parts.append(prefix + t)
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            parts.append("| " + " | ".join(cells) + " |")
    return "\n\n".join(parts)


def pick_formal_file(task_dir: Path, software_name: str = "") -> tuple[Path | None, list[str]]:
    """选出该任务唯一正式文档鉴别材料（手册 docx）；代码材料不参与手册比对。"""
    formal_dir = task_dir / "正式资料"
    if not formal_dir.is_dir():
        return None, []
    warnings: list[str] = []
    all_docx = sorted(formal_dir.glob("*.docx"))
    # 只认「文档鉴别材料」（手册）；代码材料不参与手册结构比对
    manual_docx = [c for c in all_docx if "文档鉴别材料" in c.name]
    code_docx = [c for c in all_docx if "代码" in c.name or "程序鉴别材料" in c.name]
    picked: Path | None = None
    for c in manual_docx:
        if any(pat in c.name for pat in EXCLUDE_PATTERNS):
            warnings.append(f"{c.name} 疑似模板迭代稿/临时副本，未纳入比对")
            continue
        if picked is None:
            picked = c
        else:
            warnings.append(f"{c.name} 与 {picked.name} 并存，存在多版本正式件")
    if picked is None:
        if manual_docx:
            warnings.append("文档鉴别材料仅有疑似迭代稿/副本，未纳入比对")
        elif code_docx:
            warnings.append("无文档鉴别材料正式件（仅有代码材料）")
        else:
            warnings.append("正式资料目录无 docx 文件")
    return picked, warnings


def heading_overlap(a: set[str], b: set[str]) -> float:
    denom = max(len(a), len(b))
    return len(a & b) / denom if denom else 1.0


def run_workspace(workspace: Path) -> dict[str, Any]:
    tasks: list[dict[str, Any]] = []
    for d in sorted(workspace.iterdir()):
        if not d.is_dir():
            continue
        has_material = (d / "正式资料").is_dir()
        has_draft = (d / DRAFT_MANUAL).exists()
        if not (has_material or has_draft or (d / TASK_FILE).exists()):
            continue
        name = d.name
        formal, file_warnings = pick_formal_file(d, name)
        tasks.append(
            {
                "name": name,
                "dir": d,
                "draft": d / DRAFT_MANUAL if has_draft else None,
                "formal": formal,
                "formal_text": extract_docx_text(formal) if formal else "",
                "file_warnings": file_warnings,
            }
        )
    return run_tasks(tasks)


def run_tasks(tasks: list[dict[str, Any]]) -> dict[str, Any]:
    hard_errors: list[str] = []
    risks: list[dict[str, Any]] = []

    # ── 1. 草稿手册批量结构检查（确定性 + 风险分级）──
    draft_paths = [t["draft"] for t in tasks if t["draft"]]
    if draft_paths:
        draft_report = batch_run(draft_paths, batch_id="整批")
        hard_errors.extend(draft_report["errors"])
        for r in draft_report["risks"]:
            r["source"] = "draft"
            risks.append(r)

    # ── 2. 草稿 vs 正式件版本漂移（确定性硬错误）──
    for t in tasks:
        if t["draft"] is None or not t["formal_text"]:
            continue
        d_plain = set(heading_skeleton(t["draft"])[1])
        f_plain = set()
        for line in t["formal_text"].splitlines():
            m = re.match(r"^#{1,6}\s+(.+?)\s*$", line.strip())
            if m:
                f_plain.add(re.sub(r"^\d+(?:\.\d+)*[、.\s]*", "", m.group(1).strip()))
        if not f_plain:
            hard_errors.append(f"{t['name']}: 正式件无法提取标题（可能为空文档或损坏）")
            continue
        overlap = heading_overlap(d_plain, f_plain)
        if overlap < 0.5:
            hard_errors.append(
                f"{t['name']}: 草稿与正式件版本漂移——标题重合率仅 {overlap:.0%}"
                f"（草稿 {len(d_plain)} 个标题，正式件 {len(f_plain)} 个标题）"
            )

    # ── 3. 正式件之间的结构风险 ──
    formal_tasks = [t for t in tasks if t["formal_text"]]
    if len(formal_tasks) >= 2:
        # 用临时目录写正式件文本，复用 batch_run
        import tempfile
        with tempfile.TemporaryDirectory() as tmp:
            paths = []
            for t in formal_tasks:
                p = Path(tmp) / f"{t['name']}__formal.md"
                p.write_text(t["formal_text"], encoding="utf-8")
                paths.append(p)
            formal_report = batch_run(paths, batch_id="整批-正式件")
            # 清理临时目录前缀，恢复任务名
            def _clean(p: str) -> str:
                return re.sub(r"^tmp[^/]*/", "", p).replace("__formal.md", "")
            hard_errors.extend(_clean(e) for e in formal_report["errors"])
            for r in formal_report["risks"]:
                r["pair"] = [_clean(p) for p in (r.get("pair") or [])]
                r["source"] = "formal"
                risks.append(r)

    # ── 4. 文件候选警告 ──
    file_warnings: list[str] = []
    for t in tasks:
        for w in t["file_warnings"]:
            file_warnings.append(f"{t['name']}: {w}")
        if t["draft"] is None and t["formal"] is not None:
            file_warnings.append(f"{t['name']}: 无草稿手册，仅有正式件（生成溯源缺失）")
        if t["draft"] is not None and t["formal"] is None:
            file_warnings.append(f"{t['name']}: 有草稿无正式件（可能尚未生成）")

    status = "blocked" if hard_errors else ("risk" if risks else "pass")
    return {
        "status": status,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "task_count": len(tasks),
        "tasks": [t["name"] for t in tasks],
        "hard_errors": hard_errors,
        "risks": risks,
        "file_warnings": file_warnings,
    }


def write_report(report: dict[str, Any], out_dir: Path) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / "整批复检报告.json").write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    lines = [
        "# 软著整批复检报告",
        "",
        f"- 生成时间：{report['generated_at']}",
        f"- 任务数：{report['task_count']}",
        f"- 状态：{report['status']}（blocked=有确定性硬错误；risk=有相似度风险；pass=通过）",
        "",
        "## 确定性硬错误（需修复，不涉及返工决策）",
        "",
    ]
    if report["hard_errors"]:
        for e in report["hard_errors"]:
            lines.append(f"- {e}")
    else:
        lines.append("无。")
    lines.extend(["", "## 相似度风险（高风险需人工复核，中低风险放行）", ""])
    if report["risks"]:
        for r in sorted(report["risks"], key=lambda x: {"high": 0, "medium": 1, "low": 2}.get(x["level"], 3)):
            pair = " 与 ".join(r.get("pair") or [])
            lines.append(f"- [{r['level']}]（{r.get('source','')}）{pair}：{r.get('detail','')}")
    else:
        lines.append("无。")
    lines.extend(["", "## 文件候选警告", ""])
    if report["file_warnings"]:
        for w in report["file_warnings"]:
            lines.append(f"- {w}")
    else:
        lines.append("无。")
    lines.append("")
    (out_dir / "整批复检报告.md").write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--workspace", required=True, help="<年份>年软件著作权申请资料 目录")
    parser.add_argument("--out", help="报告输出目录（默认工作区下的 整批复检报告/）")
    parser.add_argument("--json", action="store_true", help="同时打印 JSON")
    args = parser.parse_args()

    workspace = Path(args.workspace)
    if not workspace.is_dir():
        raise SystemExit(f"workspace 不存在: {workspace}")

    report = run_workspace(workspace)
    out_dir = Path(args.out) if args.out else workspace / "整批复检报告"
    write_report(report, out_dir)

    print(f"BATCH RISK REPORT {report['status'].upper()}（任务 {report['task_count']} 个）")
    print(f"  确定性硬错误: {len(report['hard_errors'])}")
    for e in report["hard_errors"][:20]:
        print(f"    E: {e}")
    from collections import Counter
    levels = Counter(r["level"] for r in report["risks"])
    print(f"  相似度风险: {dict(levels)}")
    for r in [r for r in report["risks"] if r["level"] == "high"][:15]:
        print(f"    HIGH: {' 与 '.join(r.get('pair') or [])}: {r.get('detail','')[:80]}")
    print(f"  文件候选警告: {len(report['file_warnings'])}")
    print(f"报告已写入: {out_dir}")
    if args.json:
        print(json.dumps(report, ensure_ascii=False, indent=2))
    sys.exit(0)


if __name__ == "__main__":
    main()
