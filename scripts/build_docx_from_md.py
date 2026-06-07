#!/usr/bin/env python3
"""Build final DOCX/TXT files from confirmed Markdown drafts."""

from __future__ import annotations

import argparse
import html
import re
import shutil
import subprocess
import tempfile
import time
import zipfile
from pathlib import Path
from typing import Any

from common import confirm_params, ensure_dir, read_json, resolve_workdir, safe_filename

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor

    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


BLACK_RGB = "000000"


def replace_generated_docx(tmp_path: Path, docx_path: Path) -> None:
    """Replace generated DOCX on Windows where existing files may be briefly held."""
    for attempt in range(5):
        try:
            if docx_path.exists():
                docx_path.unlink()
            tmp_path.replace(docx_path)
            return
        except PermissionError:
            if attempt == 4:
                raise
            time.sleep(0.25)


def variant_output_path(path: Path, suffix: str) -> Path:
    candidate = path.with_name(f"{path.stem}{suffix}{path.suffix}")
    if not candidate.exists():
        return candidate
    for index in range(2, 100):
        numbered = path.with_name(f"{path.stem}{suffix}{index}{path.suffix}")
        if not numbered.exists():
            return numbered
    raise RuntimeError(f"无法生成未占用的输出文件名：{path.name}")


def strip_markdown_links(text: str) -> str:
    text = re.sub(r"(?<!!)\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    text = re.sub(r"<(https?://[^>]+)>", r"\1", text)
    return text


def add_formatted_paragraph(document: Any, text: str, font_name: str = "SimSun", font_size: float = 10.5) -> None:
    """Add a paragraph with **bold** spans parsed into separate runs."""
    p = document.add_paragraph()
    add_formatted_paragraph_runs(p, text, font_name, font_size)


def add_formatted_paragraph_runs(p: Any, text: str, font_name: str = "SimSun", font_size: float = 10.5) -> None:
    """Add formatted runs (with **bold** spans) to an existing paragraph."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        set_run_font(run, font_name, font_size)


def parse_application_lines(md_path: Path) -> tuple[list[str], list[str]]:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    # Format A: ➤ lines
    fields = [line.strip() for line in lines if line.strip().startswith("➤")]
    if fields:
        warnings = [line for line in fields if "待用户确认" in line]
        return fields, warnings

    # Format B: table rows (| 序号 | 字段 | 值 | 状态 |)
    table_fields: list[str] = []
    warnings: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|") or "序号" in stripped or "---" in stripped:
            continue
        cols = [c.strip() for c in stripped.split("|")[1:-1]]
        if len(cols) >= 3:
            field_name = cols[1]
            value = cols[2]
            status = cols[3] if len(cols) > 3 else ""
            line_str = f"➤{field_name}：{value}"
            table_fields.append(line_str)
            if "待用户确认" in status or "待填写" in status or "待确认" in status:
                warnings.append(line_str)
    return table_fields, warnings


def parse_application_field(md_path: Path, field_name: str) -> str:
    if not md_path.exists():
        return ""
    prefix = f"➤{field_name}："
    for line in md_path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if stripped.startswith(prefix):
            return stripped[len(prefix) :].strip()
    return ""


def application_version(draft_dir: Path) -> str:
    version = parse_application_field(draft_dir / "申请表信息.md", "版本号")
    if "待用户确认" in version:
        return ""
    return version


def application_software_name(draft_dir: Path) -> str:
    name = parse_application_field(draft_dir / "申请表信息.md", "软件全称")
    if "待用户确认" in name:
        return ""
    return name


def write_application_txt(draft_dir: Path, out_dir: Path) -> tuple[Path | None, list[str]]:
    md_path = draft_dir / "申请表信息.md"
    if not md_path.exists():
        return None, ["缺少草稿/申请表信息.md"]
    fields, warnings = parse_application_lines(md_path)
    out_path = out_dir / "申请表信息.txt"
    out_path.write_text("\n".join(fields) + "\n", encoding="utf-8")
    return out_path, warnings


def read_json_if_exists(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    return read_json(path)


def confirmation_issues(workdir: Path) -> list[str]:
    draft_dir = workdir / "草稿"
    issues: list[str] = []

    gates = read_json_if_exists(workdir / "门禁状态.json")

    if not gates.get("business", {}).get("confirmed"):
        issues.append("业务理解尚未确认：请确认 草稿/业务理解.md 后记录 business 门禁")

    if not gates.get("code-selection", {}).get("confirmed"):
        issues.append("代码文件选择尚未确认：请确认 草稿/代码文件选择.json 后记录 code-selection 门禁")

    if not gates.get("screenshot-method", {}).get("confirmed"):
        issues.append("截图方式尚未确认：请选择截图方式后记录 screenshot-method 门禁")

    if not gates.get("application-fields", {}).get("confirmed"):
        issues.append("申请表字段尚未确认：请补全字段后记录 application-fields 门禁")

    if not gates.get("markdown", {}).get("confirmed"):
        issues.append("Markdown 草稿尚未最终确认：请确认全部草稿后记录 markdown 门禁")

    app_md = draft_dir / "申请表信息.md"
    if app_md.exists():
        _, warnings = parse_application_lines(app_md)
        if warnings:
            issues.append("申请表信息仍包含\"待用户确认\"字段")
    else:
        issues.append("缺少 草稿/申请表信息.md")

    return issues


def parse_code_pages(md_path: Path) -> list[tuple[int, list[str]]]:
    pages: list[tuple[int, list[str]]] = []
    current_no: int | None = None
    current_lines: list[str] = []
    in_fence = False

    for raw in md_path.read_text(encoding="utf-8").splitlines():
        page_match = re.match(r"^##\s+第\s*(\d+)\s*页", raw.strip())
        if page_match:
            if current_no is not None:
                pages.append((current_no, current_lines))
            current_no = int(page_match.group(1))
            current_lines = []
            in_fence = False
            continue
        if raw.strip().startswith("```"):
            in_fence = not in_fence
            continue
        if current_no is not None and in_fence:
            current_lines.append(raw)

    if current_no is not None:
        pages.append((current_no, current_lines))
    return pages


def set_run_font(run: Any, name: str, size_pt: float) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    try:
        run.font.color.rgb = RGBColor(0, 0, 0)
    except Exception:
        pass
    try:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), name)
        rFonts.set(qn("w:hAnsi"), name)
        rFonts.set(qn("w:eastAsia"), name)
    except Exception:
        pass


def set_normal_font(document: Any, name: str = "SimSun", size_pt: float = 10.5) -> None:
    style = document.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)
    try:
        style.font.color.rgb = RGBColor(0, 0, 0)
    except Exception:
        pass
    try:
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), name)
        rFonts.set(qn("w:hAnsi"), name)
        rFonts.set(qn("w:eastAsia"), name)
    except Exception:
        pass


def set_style_black(document: Any) -> None:
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        try:
            style = document.styles[style_name]
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.name = "SimSun"
            # Force eastAsia font in the style definition itself
            rPr = style.element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:ascii"), "SimSun")
            rFonts.set(qn("w:hAnsi"), "SimSun")
            rFonts.set(qn("w:eastAsia"), "SimSun")
        except Exception:
            pass


def force_black_document(document: Any) -> None:
    set_style_black(document)
    containers = [document]
    for section in document.sections:
        containers.extend([section.header, section.footer])
    for container in containers:
        for paragraph in container.paragraphs:
            for run in paragraph.runs:
                try:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                except Exception:
                    pass
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            try:
                                run.font.color.rgb = RGBColor(0, 0, 0)
                            except Exception:
                                pass


def configure_a4(document: Any) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)


def configure_code_a4(document: Any) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)


def add_page_field(paragraph: Any) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for element in (begin, instr, separate, result, end):
        run = paragraph.add_run()
        run._r.append(element)
        set_run_font(run, "SimSun", 8)


def set_code_header(document: Any, software_name: str, version: str) -> None:
    section = document.sections[0]
    section.header.is_linked_to_previous = False
    header = section.header
    header.paragraphs[0].text = "" if header.paragraphs else None

    # Build a two-column header: software name on the left, page number on the right.
    table = header.add_table(rows=1, cols=2, width=Cm(17.5))
    table.autofit = True
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]

    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_para.paragraph_format.space_before = Pt(0)
    left_para.paragraph_format.space_after = Pt(0)
    left_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    left_para.paragraph_format.line_spacing = Pt(12)
    left_run = left_para.add_run(f"{software_name} {version}")
    set_run_font(left_run, "SimSun", 8)

    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_para.paragraph_format.space_before = Pt(0)
    right_para.paragraph_format.space_after = Pt(0)
    right_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    right_para.paragraph_format.line_spacing = Pt(12)
    prefix = right_para.add_run("第 ")
    set_run_font(prefix, "SimSun", 8)
    add_page_field(right_para)
    suffix = right_para.add_run(" 页")
    set_run_font(suffix, "SimSun", 8)

    # Remove borders from the header table
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = OxmlElement("w:tcBorders")
        for border_name in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "nil")
            tc_borders.append(border)
        tc_pr.append(tc_borders)


def build_code_docx_python(md_path: Path, out_path: Path, software_name: str, version: str) -> None:
    pages = parse_code_pages(md_path)
    if not pages:
        raise RuntimeError(f"No code pages parsed from {md_path}")

    document = Document()
    configure_code_a4(document)
    set_normal_font(document, "Consolas", 7.2)
    set_style_black(document)
    set_code_header(document, software_name, version)

    for index, (page_no, lines) in enumerate(pages):
        for line in lines:
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(12)
            run = p.add_run(line if line else " ")
            set_run_font(run, "Consolas", 7.2)
        if index != len(pages) - 1:
            document.add_page_break()

    force_black_document(document)
    document.save(out_path)


def paragraph_xml(text: str, font: str = "SimSun", size_half_points: int = 21, align: str | None = None, line_twips: int = 240) -> str:
    align_xml = f'<w:jc w:val="{align}"/>' if align else ""
    escaped = html.escape(text)
    return (
        "<w:p>"
        f"<w:pPr>{align_xml}<w:spacing w:after=\"0\" w:line=\"{line_twips}\" w:lineRule=\"exact\"/></w:pPr>"
        "<w:r>"
        f"<w:rPr><w:rFonts w:ascii=\"{font}\" w:hAnsi=\"{font}\" w:eastAsia=\"{font}\"/>"
        f"<w:color w:val=\"{BLACK_RGB}\"/>"
        f"<w:sz w:val=\"{size_half_points}\"/><w:szCs w:val=\"{size_half_points}\"/></w:rPr>"
        f"<w:t xml:space=\"preserve\">{escaped}</w:t>"
        "</w:r>"
        "</w:p>"
    )


def page_break_xml() -> str:
    return '<w:p><w:r><w:br w:type="page"/></w:r></w:p>'


def page_field_runs_xml() -> str:
    return (
        '<w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/>'
        f'<w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
        '<w:fldChar w:fldCharType="begin"/></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/>'
        f'<w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
        '<w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/>'
        f'<w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
        '<w:fldChar w:fldCharType="separate"/></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/>'
        f'<w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
        '<w:t>1</w:t></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/>'
        f'<w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
        '<w:fldChar w:fldCharType="end"/></w:r>'
    )


def header_xml(header_text: str) -> str:
    """Build a two-column header: software name left, page number right."""
    escaped = html.escape(header_text)
    # Use a borderless table for left/right alignment in header
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:hdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:tbl>
    <w:tblPr>
      <w:tblW w:w="5000" w:type="pct"/>
      <w:tblBorders>
        <w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/><w:insideH w:val="nil"/><w:insideV w:val="nil"/>
      </w:tblBorders>
    </w:tblPr>
    <w:tr>
      <w:tc>
        <w:p>
          <w:pPr><w:jc w:val="left"/><w:spacing w:after="0" w:line="240" w:lineRule="exact"/></w:pPr>
          <w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/><w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr><w:t xml:space="preserve">{escaped}</w:t></w:r>
        </w:p>
      </w:tc>
      <w:tc>
        <w:p>
          <w:pPr><w:jc w:val="right"/><w:spacing w:after="0" w:line="240" w:lineRule="exact"/></w:pPr>
          <w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/><w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr><w:t xml:space="preserve">第 </w:t></w:r>
          {page_field_runs_xml()}
          <w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/><w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr><w:t xml:space="preserve"> 页</w:t></w:r>
        </w:p>
      </w:tc>
    </w:tr>
  </w:tbl>
</w:hdr>"""


def minimal_docx(out_path: Path, body_xml: str, header_text: str | None = None) -> None:
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
  <Override PartName="/word/header1.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"/>
</Types>"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
    header_rel = (
        '<Relationship Id="rIdHeader1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/header" Target="header1.xml"/>'
        if header_text
        else ""
    )
    doc_rels = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">{header_rel}</Relationships>"""
    styles = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal">
    <w:name w:val="Normal"/>
    <w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/><w:color w:val="{BLACK_RGB}"/><w:sz w:val="21"/></w:rPr>
  </w:style>
</w:styles>"""
    document = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:body>
    {body_xml}
      <w:sectPr>
        {'<w:headerReference w:type="default" r:id="rIdHeader1"/>' if header_text else ''}
        <w:pgSz w:w="11906" w:h="16838"/>
      <w:pgMar w:top="1134" w:right="1134" w:bottom="1134" w:left="1418" w:header="283" w:footer="283" w:gutter="0"/>
    </w:sectPr>
  </w:body>
</w:document>"""
    with zipfile.ZipFile(out_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/_rels/document.xml.rels", doc_rels)
        zf.writestr("word/styles.xml", styles)
        zf.writestr("word/document.xml", document)
        if header_text:
            zf.writestr("word/header1.xml", header_xml(header_text))


def force_black_xml(xml: str) -> str:
    xml = re.sub(r"<w:hyperlink\b[^>]*>", "", xml)
    xml = xml.replace("</w:hyperlink>", "")
    xml = re.sub(r"<w:color\b[^>]*/>", f'<w:color w:val="{BLACK_RGB}"/>', xml)

    def ensure_rpr_color(match: re.Match[str]) -> str:
        value = match.group(0)
        if "<w:color" in value:
            return value
        return value.replace("</w:rPr>", f'<w:color w:val="{BLACK_RGB}"/></w:rPr>')

    xml = re.sub(r"<w:rPr\b[^>]*>.*?</w:rPr>", ensure_rpr_color, xml, flags=re.S)
    xml = re.sub(r"<w:r>(?!<w:rPr>)", f'<w:r><w:rPr><w:color w:val="{BLACK_RGB}"/></w:rPr>', xml)
    return xml


def normalize_docx_text_color(docx_path: Path) -> None:
    tmp_path = docx_path.with_suffix(docx_path.suffix + ".tmp")
    color_xml_parts = (
        "word/document.xml",
        "word/styles.xml",
        "word/numbering.xml",
        "word/header",
        "word/footer",
    )
    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename.endswith(".xml") and item.filename.startswith(color_xml_parts):
                text = data.decode("utf-8")
                data = force_black_xml(text).encode("utf-8")
            elif item.filename.endswith(".rels"):
                text = data.decode("utf-8", errors="ignore")
                if "hyperlink" in text:
                    text = re.sub(r'\s*<Relationship\b[^>]*Type="[^"]*/hyperlink"[^>]*/>', "", text)
                    data = text.encode("utf-8")
            dst.writestr(item, data)
    replace_generated_docx(tmp_path, docx_path)


def next_header_part(names: set[str]) -> tuple[str, str]:
    index = 1
    while f"word/header{index}.xml" in names:
        index += 1
    return f"word/header{index}.xml", f"header{index}.xml"


def unique_relationship_id(rels_xml: str, base: str = "rIdManualHeader") -> str:
    if f'Id="{base}"' not in rels_xml:
        return base
    index = 2
    while f'Id="{base}{index}"' in rels_xml:
        index += 1
    return f"{base}{index}"


def add_header_to_existing_docx(docx_path: Path, header_text: str) -> None:
    """Add the same two-column header used by code materials to an existing DOCX."""
    tmp_path = docx_path.with_suffix(docx_path.suffix + ".tmp")
    with zipfile.ZipFile(docx_path, "r") as src:
        names = set(src.namelist())
        header_part, header_target = next_header_part(names)
        rels_xml = src.read("word/_rels/document.xml.rels").decode("utf-8")
        rel_id = unique_relationship_id(rels_xml)

        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as dst:
            for item in src.infolist():
                data = src.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    text = data.decode("utf-8")
                    override = (
                        f'<Override PartName="/{header_part}" '
                        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"/>'
                    )
                    if f'PartName="/{header_part}"' not in text:
                        text = text.replace("</Types>", f"{override}</Types>")
                    data = text.encode("utf-8")
                elif item.filename == "word/_rels/document.xml.rels":
                    text = data.decode("utf-8")
                    relationship = (
                        f'<Relationship Id="{rel_id}" '
                        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/header" '
                        f'Target="{header_target}"/>'
                    )
                    text = text.replace("</Relationships>", f"{relationship}</Relationships>")
                    data = text.encode("utf-8")
                elif item.filename == "word/document.xml":
                    text = data.decode("utf-8")
                    if "xmlns:r=" not in text:
                        text = text.replace(
                            "<w:document ",
                            '<w:document xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" ',
                            1,
                        )
                    header_ref = f'<w:headerReference w:type="default" r:id="{rel_id}"/>'
                    if "<w:headerReference" in text:
                        text = re.sub(r"<w:headerReference\b[^>]*/>", header_ref, text, count=1)
                    else:
                        text = re.sub(r"(<w:sectPr\b[^>]*>)", rf"\1{header_ref}", text, count=1)
                    data = text.encode("utf-8")
                dst.writestr(item, data)
            dst.writestr(header_part, header_xml(header_text))
    replace_generated_docx(tmp_path, docx_path)


def build_code_docx_ooxml(md_path: Path, out_path: Path, software_name: str, version: str) -> None:
    pages = parse_code_pages(md_path)
    if not pages:
        raise RuntimeError(f"No code pages parsed from {md_path}")
    body: list[str] = []
    for index, (page_no, lines) in enumerate(pages):
        for line in lines:
            body.append(paragraph_xml(line if line else " ", font="Consolas", size_half_points=14, line_twips=240))
        if index != len(pages) - 1:
            body.append(page_break_xml())
    minimal_docx(out_path, "\n".join(body), header_text=f"{software_name} {version}")


def add_markdown_table(document: Any, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = len(rows[0])
    cleaned = [[strip_markdown_links(c) for c in r] for r in rows]

    # 1 — create table and fill content
    table = document.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    for idx, text in enumerate(cleaned[0]):
        table.rows[0].cells[idx].text = text
    for row in cleaned[1:]:
        cells = table.add_row().cells
        for idx, text in enumerate(row[:col_count]):
            cells[idx].text = text

    # 2 — measure longest content per column (CJK = 2.0, Latin = 1.0)
    length_units = [0.0] * col_count
    for row in cleaned:
        for i, text in enumerate(row[:col_count]):
            length = 0.0
            for ch in str(text):
                if '一' <= ch <= '鿿':
                    length += 2.0
                else:
                    length += 1.0
            if length > length_units[i]:
                length_units[i] = length

    # 3 — proportional allocation → scale to A4 → then clamp minimum
    A4_CM = 15.29
    MIN_CM = 1.2
    total_units = sum(length_units) or 1.0
    raw_cm = [(lu / total_units) * A4_CM for lu in length_units]
    raw_total = sum(raw_cm)
    if raw_total > A4_CM:
        scale = A4_CM / raw_total
        raw_cm = [w * scale for w in raw_cm]

    # Apply minimum only after scaling — then re-scale if total exceeds page
    clamped = [max(w, MIN_CM) for w in raw_cm]
    clamped_total = sum(clamped)
    if clamped_total > A4_CM:
        scale2 = A4_CM / clamped_total
        clamped = [w * scale2 for w in clamped]

    # 4 — set fixed layout, write column and cell widths
    table.allow_autofit = False
    table.autofit = False

    for col_idx in range(col_count):
        width = Cm(clamped[col_idx])
        table.columns[col_idx].width = width
        for row in table.rows:
            row.cells[col_idx].width = width


def parse_table_line(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_image(document: Any, image_path: Path) -> None:
    if not image_path.exists():
        p = document.add_paragraph()
        run = p.add_run(f"[截图缺失：{image_path}]")
        set_run_font(run, "SimSun", 10.5)
        return
    try:
        document.add_picture(str(image_path), width=Inches(5.8))
    except Exception:
        p = document.add_paragraph()
        run = p.add_run(f"[截图无法插入：{image_path}]")
        set_run_font(run, "SimSun", 10.5)


def add_toc_field(paragraph: Any, instruction: str = ' TOC \\o "1-3" \\h \\z \\u ') -> None:
    """Insert a Word TOC field that users can right-click to update."""
    hint = '右键点击此处选择"更新域"以生成目录'
    for ch, text in [("begin", None), (None, instruction), ("separate", None), (None, hint), ("end", None)]:
        run = paragraph.add_run()
        if ch:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), ch)
            run._r.append(fld)
        if text:
            instr = OxmlElement("w:instrText")
            instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            instr.text = text
            run._r.append(instr)
        set_run_font(run, "SimSun", 10.5)


def build_manual_docx_python(md_path: Path, out_path: Path, base_dir: Path, software_name: str, version: str) -> None:
    document = Document()
    configure_a4(document)
    set_normal_font(document, "SimSun", 10.5)
    set_style_black(document)
    set_code_header(document, software_name, version)

    # ── Cover page: vertically + horizontally centered ──
    # Set section vertical alignment to center
    sectPr = document.sections[0]._sectPr
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    sectPr.append(vAlign)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(software_name)
    set_run_font(run, "SimSun", 26)
    run.bold = True

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("用户使用说明书")
    set_run_font(run, "SimSun", 26)
    run.bold = True

    document.add_page_break()

    # ── Body section: restore normal vertical alignment (top) ──
    body_section = document.add_section(WD_SECTION.NEW_PAGE)
    body_section.top_margin = Cm(2.54)
    body_section.bottom_margin = Cm(2.54)
    body_section.left_margin = Cm(3.17)
    body_section.right_margin = Cm(2.54)
    body_sectPr = body_section._sectPr
    body_vAlign = OxmlElement("w:vAlign")
    body_vAlign.set(qn("w:val"), "top")
    body_sectPr.append(body_vAlign)

    # ── Table of contents ──
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("目  录")
    set_run_font(run, "SimSun", 20)
    run.bold = True

    toc_para = document.add_paragraph()
    toc_para.paragraph_format.space_before = Pt(0)
    toc_para.paragraph_format.space_after = Pt(0)
    add_toc_field(toc_para)

    document.add_page_break()

    # ── Body: skip Markdown preamble and parse from first content heading ──
    raw_lines = md_path.read_text(encoding="utf-8").splitlines()
    lines: list[str] = []
    started = False
    for line in raw_lines:
        if not started and (re.match(r"^##\s+\d+\.\s+", line.strip()) or re.match(r"^##\s+[一二三四五六七八九十]+、", line.strip())):
            started = True
        if started:
            lines.append(line)
    table_buf: list[list[str]] = []
    in_fence = False

    def flush_table() -> None:
        nonlocal table_buf
        if table_buf:
            data = [row for row in table_buf if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in row)]
            add_markdown_table(document, data)
            table_buf = []

    for line in lines:
        stripped = line.strip()
        stripped = strip_markdown_links(stripped)
        if stripped.startswith("```"):
            flush_table()
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        if re.match(r"> \*\*模块类型：", stripped):
            continue
        if stripped.startswith("<!--") and "截图" in stripped:
            stripped = "【截图预留：请在此处插入当前功能页面或操作结果截图。】"
        if stripped.startswith("|") and stripped.endswith("|"):
            table_buf.append(parse_table_line(stripped))
            continue
        flush_table()
        if not stripped:
            continue
        if stripped == r"\newpage":
            document.add_page_break()
            continue
        image_match = re.search(r"!\[[^\]]*\]\(([^)]+)\)", stripped)
        if image_match:
            add_image(document, (base_dir / image_match.group(1)).resolve())
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 3)
            p = document.add_heading(heading.group(2), level=level)
            for run in p.runs:
                try:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    rPr = run._r.get_or_add_rPr()
                    rFonts = rPr.find(qn("w:rFonts"))
                    if rFonts is None:
                        rFonts = OxmlElement("w:rFonts")
                        rPr.insert(0, rFonts)
                    rFonts.set(qn("w:ascii"), "SimSun")
                    rFonts.set(qn("w:hAnsi"), "SimSun")
                    rFonts.set(qn("w:eastAsia"), "SimSun")
                except Exception:
                    pass
            continue
        if re.match(r"^[-*+]\s+", stripped):
            p = document.add_paragraph(style="List Bullet")
            text = re.sub(r"^[-*+]\s+", "", stripped)
            add_formatted_paragraph_runs(p, text)
            continue
        if re.match(r"^\d+\.\s+", stripped):
            p = document.add_paragraph(style="List Number")
            text = re.sub(r"^\d+\.\s+", "", stripped)
            add_formatted_paragraph_runs(p, text)
            continue
        add_formatted_paragraph(document, stripped, "SimSun", 10.5)
    flush_table()
    force_black_document(document)
    document.save(out_path)


def pandoc_available() -> bool:
    return shutil.which("pandoc") is not None


def build_with_pandoc(md_path: Path, out_path: Path, code_mode: bool = False) -> None:
    if not pandoc_available():
        raise RuntimeError("python-docx is unavailable and pandoc is not installed")
    source = md_path
    tmp_name: str | None = None
    original_text = md_path.read_text(encoding="utf-8")
    text = original_text
    text = re.sub(r"```text\s*\nSTOP_FOR_USER\n.*?```", "", text, flags=re.S)
    text = re.sub(r"<!--[^>]*截图[^>]*-->", "【截图预留：请在此处插入当前功能页面或操作结果截图。】", text)
    text = strip_markdown_links(text)
    if code_mode:
        text = re.sub(r"(?=^##\s+第\s*\d+\s*页)", r"\n\\newpage\n", text, flags=re.M)
    if code_mode or "STOP_FOR_USER" in original_text:
        with tempfile.NamedTemporaryFile("w", suffix=".md", delete=False, encoding="utf-8") as tmp:
            tmp.write(text)
            tmp_name = tmp.name
        source = Path(tmp_name)
    try:
        subprocess.run(["pandoc", "-f", "markdown", "-t", "docx", str(source), "-o", str(out_path)], check=True)
    finally:
        if tmp_name:
            Path(tmp_name).unlink(missing_ok=True)


def build_code_docx(md_path: Path, out_path: Path, software_name: str, version: str) -> None:
    if DOCX_AVAILABLE:
        build_code_docx_python(md_path, out_path, software_name, version)
    else:
        build_code_docx_ooxml(md_path, out_path, software_name, version)
    normalize_docx_text_color(out_path)


def build_manual_docx(md_path: Path, out_path: Path, base_dir: Path, software_name: str, version: str) -> None:
    if DOCX_AVAILABLE:
        build_manual_docx_python(md_path, out_path, base_dir, software_name, version)
    else:
        build_with_pandoc(md_path, out_path, code_mode=False)
        add_header_to_existing_docx(out_path, f"{software_name} {version}")
    normalize_docx_text_color(out_path)


def run_command(command: list[str], cwd: Path | None = None, timeout: int = 60) -> tuple[int, str]:
    try:
        completed = subprocess.run(command, cwd=cwd, text=True, capture_output=True, timeout=timeout)
        return completed.returncode, (completed.stdout + completed.stderr).strip()
    except Exception as exc:
        return 99, str(exc)


def find_dotnet() -> str | None:
    found = shutil.which("dotnet")
    if found:
        return found
    default = Path(r"C:\Program Files\dotnet\dotnet.exe")
    if default.exists():
        return str(default)
    return None


def find_docx_toolkit_cli(skill_dir: Path) -> Path | None:
    bin_dir = skill_dir / "vendor/docx-toolkit/scripts/dotnet/DocxToolkit.Cli/bin"
    candidates = sorted(bin_dir.glob("**/DocxToolkit.Cli.dll"), reverse=True)
    return candidates[0] if candidates else None


def preview_docx_text(docx_path: Path, max_chars: int = 500) -> str:
    try:
        with zipfile.ZipFile(docx_path, "r") as docx:
            xml = docx.read("word/document.xml").decode("utf-8", errors="ignore")
    except Exception as exc:
        return f"Preview failed: {exc}"
    xml = re.sub(r"<w:tab\s*/>", "\t", xml)
    xml = re.sub(r"</w:p>", "\n", xml)
    text = re.sub(r"<[^>]+>", "", xml)
    text = html.unescape(text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text[:max_chars]


def docx_checks(skill_dir: Path, outputs: list[Path]) -> list[str]:
    notes: list[str] = []
    dotnet = find_dotnet()
    cli = find_docx_toolkit_cli(skill_dir)
    if dotnet and cli:
        code, output = run_command([dotnet, "--version"], timeout=15)
        status = "READY" if code == 0 else "NOT READY"
        notes.append(f"DOCX env: {status}\n\n```text\n.NET SDK {output}\nDocxToolkit.Cli {cli}\n```")
        for out in outputs:
            code, output = run_command([dotnet, str(cli), "validate", "--input", str(out), "--json"], timeout=45)
            first_lines = "\n".join(output.splitlines()[:12])
            notes.append(f"Validate {out.name}: exit={code}\n\n```text\n{first_lines}\n```")
            preview = preview_docx_text(out)
            notes.append(f"Preview {out.name}\n\n```text\n{preview}\n```")
        return notes

    env_script = skill_dir / "vendor/docx-toolkit/scripts/env_check.sh"
    preview_script = skill_dir / "vendor/docx-toolkit/scripts/docx_preview.sh"
    if env_script.exists():
        code, output = run_command(["bash", str(env_script)], cwd=env_script.parent.parent, timeout=30)
        status = "READY" if code == 0 else "NOT READY"
        first_lines = "\n".join(output.splitlines()[:12])
        notes.append(f"DOCX env: {status}\n\n```text\n{first_lines}\n```")
    else:
        notes.append("DOCX env: vendor script missing")

    if preview_script.exists():
        for out in outputs:
            code, output = run_command(["bash", str(preview_script), str(out)], timeout=45)
            first_lines = "\n".join(output.splitlines()[:8])
            notes.append(f"Preview {out.name}: exit={code}\n\n```text\n{first_lines}\n```")
    return notes


def build_all(workdir: Path, software_name: str, version: str, skip_preview: bool) -> dict[str, Any]:
    workdir = ensure_dir(workdir)
    draft_dir = workdir / "草稿"
    final_dir = ensure_dir(workdir / "正式资料")
    app_name = application_software_name(draft_dir)
    app_version = application_version(draft_dir)
    final_software_name = app_name or software_name
    final_version = app_version or version
    safe_name = safe_filename(final_software_name)
    outputs: list[Path] = []
    warnings: list[str] = []
    if app_name and app_name != software_name:
        warnings.append(f"命令参数软件名称为 {software_name}，正式资料已按申请表信息软件名称 {app_name} 生成")
    if app_version and app_version != version:
        warnings.append(f"命令参数版本号为 {version}，正式资料已按申请表信息版本号 {app_version} 生成")
    screenshot_confirmation = read_json_if_exists(workdir / "门禁状态.json")
    screenshot_method = screenshot_confirmation.get("screenshot-method", {}).get("method")
    screenshot_manifest = workdir / "截图/截图清单.json"
    if screenshot_method == "skip":
        warnings.append("技术图表已嵌入——用户选择暂不截图的页面截图已保留占位符，存在补正风险")
    elif screenshot_method and not screenshot_manifest.exists():
        warnings.append("技术图表已嵌入——用户页面截图尚未提供或未运行截图整理，操作手册中截图位置为占位符，存在补正风险")
    elif screenshot_manifest.exists():
        screenshots = read_json_if_exists(screenshot_manifest).get("screenshots") or []
        if not screenshots:
            warnings.append("操作手册截图清单为空；操作手册应保留截图预留位置")

    app_txt, app_warnings = write_application_txt(draft_dir, final_dir)
    if app_txt:
        outputs.append(app_txt)
    warnings.extend(app_warnings)

    code_specs = [
        ("代码-前30页.md", f"{safe_name}-代码(前30页).docx"),
        ("代码-后30页.md", f"{safe_name}-代码(后30页).docx"),
        ("代码-全部.md", f"{safe_name}-代码(全部).docx"),
    ]
    for md_name, docx_name in code_specs:
        md_path = draft_dir / md_name
        if md_path.exists():
            out_path = final_dir / docx_name
            build_code_docx(md_path, out_path, final_software_name, final_version)
            outputs.append(out_path)

    manual_md = draft_dir / "操作手册.md"
    if manual_md.exists():
        manual_out = final_dir / f"{safe_name}_操作手册.docx"
        manual_source = manual_md
        tmp_manual: Path | None = None
        if app_name and app_name != software_name:
            text = manual_md.read_text(encoding="utf-8").replace(software_name, app_name)
            with tempfile.NamedTemporaryFile("w", suffix=".md", delete=False, encoding="utf-8") as tmp:
                tmp.write(text)
                tmp_manual = Path(tmp.name)
            manual_source = tmp_manual
        try:
            try:
                build_manual_docx(manual_source, manual_out, draft_dir.parent, final_software_name, final_version)
                outputs.append(manual_out)
            except PermissionError:
                fallback_out = variant_output_path(manual_out, "_新版模板式")
                build_manual_docx(manual_source, fallback_out, draft_dir.parent, final_software_name, final_version)
                outputs.append(fallback_out)
                warnings.append(f"原操作手册文件被占用，已生成备用正式文件：{fallback_out.name}")
        finally:
            if tmp_manual:
                tmp_manual.unlink(missing_ok=True)
    else:
        warnings.append("缺少草稿/操作手册.md")

    skill_dir = Path(__file__).resolve().parents[1]
    notes = [] if skip_preview else docx_checks(skill_dir, [p for p in outputs if p.suffix.lower() == ".docx"])
    report = write_report(final_dir, outputs, warnings, notes)
    return {"outputs": [str(p) for p in outputs], "warnings": warnings, "report": str(report)}


def write_report(workdir: Path, outputs: list[Path], warnings: list[str], notes: list[str]) -> Path:
    report = workdir / "生成报告.md"
    lines = ["# 生成报告", "", "## 输出文件", ""]
    for path in outputs:
        size = path.stat().st_size if path.exists() else 0
        lines.append(f"- `{path.name}` ({size} bytes)")
    lines.extend(["", "## 警告", ""])
    if warnings:
        lines.extend(f"- {warning}" for warning in warnings)
    else:
        lines.append("- 无")
    lines.extend(["", "## DOCX 校验", ""])
    if notes:
        lines.extend(notes)
    else:
        lines.append("- 已跳过预览校验")
    report.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return report


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--workdir", help="Task workdir; auto-derived from --task-dir if omitted")
    parser.add_argument("--task-dir", help="Task root dir; auto-resolved from current directory if omitted")
    parser.add_argument("--software-name", required=True)
    parser.add_argument("--version", required=True)
    parser.add_argument("--skip-preview", action="store_true")
    parser.add_argument("--confirm", action="store_true", help="Confirmed by user, proceed with execution")
    args = parser.parse_args()

    workdir = Path(args.workdir) if args.workdir else resolve_workdir(args.task_dir)

    confirm_params({"输出目录": str(workdir), "软件名称": args.software_name, "版本号": args.version}, args.confirm)
    issues = confirmation_issues(workdir)
    if issues:
        print("STOP_FOR_USER")
        print("NEXT_ACTION: 正式 Word/TXT 生成前必须完成以下确认：")
        for issue in issues:
            print(f"- {issue}")
        raise SystemExit(2)

    result = build_all(workdir, args.software_name, args.version, args.skip_preview)
    print(f"OK final materials: {workdir / '正式资料'}")
    for output in result["outputs"]:
        print(output)
    if result["warnings"]:
        print("Warnings:")
        for warning in result["warnings"]:
            print(f"- {warning}")
    print(f"Report: {result['report']}")


if __name__ == "__main__":
    main()
