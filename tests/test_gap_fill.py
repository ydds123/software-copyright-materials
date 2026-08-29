"""补漏项测试：pHash 去重、申报前置阻断、跨材料一致性、最终件复检、profile、开关。"""

import json
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

SCRIPTS = Path(__file__).resolve().parents[1] / "scripts"
sys.path.insert(0, str(SCRIPTS))

import cross_material_check as cmc  # noqa: E402
import final_artifact_check as fac  # noqa: E402
import visual_evidence_check as vec  # noqa: E402
import submission_readiness_check as src  # noqa: E402


def _mock_assess(has_real_data=True):
    return patch(
        "visual_evidence_check.assess_image_sync",
        return_value={
            "ok": True,
            "cached": False,
            "model": "m",
            "model_version": "m",
            "prompt_version": "p",
            "assessed_at": "t",
            "page_type": "list",
            "has_real_data": has_real_data,
            "suspected_design_mockup": False,
            "empty_state_detected": False,
            "status_label_count": None,
            "raw_answer": "{}",
        },
    )


def _img(path: Path, color: tuple[int, int, int], size=(64, 64)) -> Path:
    from PIL import Image

    Image.new("RGB", size, color).save(path)
    return path


def _plan(visual_evidence, features=None):
    return {
        "schema_version": 3,
        "software_scope": {"name": "巡检系统", "version": "V1.0"},
        "visual_evidence": visual_evidence,
        "features": features
        or [
            {
                "feature_id": "F-001",
                "name": "核心功能",
                "importance": "core",
                "visual_evidence": [],
                "visual_status": "required",
            }
        ],
        "code_evidence": [],
    }


def _ev(eid, fpath, grade="A", capture="live_system", scrubbed=True, acquired=True):
    return {
        "evidence_id": eid,
        "kind": "ui_screenshot",
        "file_path": fpath,
        "grade": grade,
        "capture_source": capture,
        "data_state": "populated",
        "privacy_scrubbed": scrubbed,
        "acquisition_status": "acquired" if acquired else "pending",
        "mapped_features": ["F-001"],
    }


class PerceptualDupTest(unittest.TestCase):
    def test_identical_bytes_dup(self):
        with tempfile.TemporaryDirectory() as td:
            d = Path(td)
            _img(d / "a.png", (255, 0, 0))
            _img(d / "b.png", (255, 0, 0))
            plan = _plan([_ev("V-001", "a.png"), _ev("V-002", "b.png")])
            pp = d / "材料证据计划.json"
            pp.write_text(json.dumps(plan, ensure_ascii=False), encoding="utf-8")
            with _mock_assess():
                report = vec.run(pp, d, cache_path=d / "c.json")
            self.assertTrue(any("重复" in e for e in report["errors"]))

    def test_slightly_modified_still_dup(self):
        """加水印级轻微修改（改少数像素）仍被判近似重复。"""
        with tempfile.TemporaryDirectory() as td:
            d = Path(td)
            _img(d / "a.png", (200, 30, 30))
            _img(d / "b.png", (201, 31, 31))  # 极轻微色差
            h1 = vec.perceptual_hash(d / "a.png")
            h2 = vec.perceptual_hash(d / "b.png")
            self.assertIsNotNone(h1)
            self.assertIsNotNone(h2)
            dist = vec.hamming_distance(h1, h2)
            self.assertLessEqual(dist, vec.DHASH_DUP_THRESHOLD, f"dHash 距离 {dist} 超出阈值")

    def test_distinct_images_not_dup(self):
        from PIL import Image

        with tempfile.TemporaryDirectory() as td:
            d = Path(td)
            # 左黑右白渐变 vs 左白右黑渐变（结构相反）
            g1 = Image.new("L", (64, 64))
            g2 = Image.new("L", (64, 64))
            for y in range(64):
                for x in range(64):
                    g1.putpixel((x, y), x * 4)
                    g2.putpixel((x, y), 255 - x * 4)
            g1.save(d / "a.png")
            g2.save(d / "b.png")
            dist = vec.hamming_distance(
                vec.perceptual_hash(d / "a.png"), vec.perceptual_hash(d / "b.png")
            )
            self.assertGreater(dist, vec.DHASH_DUP_THRESHOLD, f"dHash 距离 {dist} 应超出阈值")


class DeclarationFirstTest(unittest.TestCase):
    def test_all_pending_blocked(self):
        with tempfile.TemporaryDirectory() as td:
            d = Path(td)
            _img(d / "a.png", (10, 10, 10))
            plan = _plan([_ev("V-001", "a.png", acquired=False)])
            plan["features"][0]["visual_evidence"] = ["V-001"]
            pp = d / "材料证据计划.json"
            pp.write_text(json.dumps(plan, ensure_ascii=False), encoding="utf-8")
            with _mock_assess():
                report = vec.run(pp, d, cache_path=d / "c.json")
            self.assertTrue(any(e.startswith("V7") for e in report["errors"]))

    def test_not_applicable_excluded_from_denominator(self):
        with tempfile.TemporaryDirectory() as td:
            d = Path(td)
            plan = _plan([])
            plan["features"][0]["visual_status"] = "not_applicable"
            pp = d / "材料证据计划.json"
            pp.write_text(json.dumps(plan, ensure_ascii=False), encoding="utf-8")
            with _mock_assess():
                report = vec.run(pp, d, cache_path=d / "c.json")
            self.assertFalse(any(e.startswith("V7") for e in report["errors"]))


class CrossMaterialTest(unittest.TestCase):
    def _plan(self, features):
        return {
            "schema_version": 3,
            "software_scope": {"name": "巡检系统", "version": "V1.0"},
            "features": features,
            "code_evidence": [],
        }

    def _write(self, plan, manual=None, app=None, manifest=None):
        tmp = tempfile.TemporaryDirectory()
        self.addCleanup(tmp.cleanup)
        d = Path(tmp.name)
        pp = d / "材料证据计划.json"
        pp.write_text(json.dumps(plan, ensure_ascii=False), encoding="utf-8")
        manual_p = d / "m.md"
        if manual is not None:
            manual_p.write_text(manual, encoding="utf-8")
        app_p = d / "申请表信息.md"
        if app is not None:
            app_p.write_text(app, encoding="utf-8")
        manifest_p = d / "代码提取清单.json"
        if manifest is not None:
            manifest_p.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
        return pp, manual_p, app_p, manifest_p

    def test_consistent_pass(self):
        plan = self._plan(
            [{"feature_id": "F-001", "name": "巡检计划管理", "importance": "core", "document_sections": []}]
        )
        pp, mp, _, _ = self._write(
            plan,
            manual="# 手册\n\n软件名称：巡检系统\n\n版本号：V1.0\n\n巡检计划管理功能说明。\n",
        )
        report = cmc.run(pp, mp)
        self.assertEqual(report["status"], "pass", report["errors"])

    def test_name_mismatch_blocked(self):
        plan = self._plan([])
        pp, mp, _, _ = self._write(plan, manual="# 手册\n\n软件名称：另一个系统\n\n版本号：V1.0\n")
        report = cmc.run(pp, mp)
        self.assertTrue(any("软件名称" in e for e in report["errors"]))

    def test_core_feature_missing_blocked(self):
        plan = self._plan(
            [{"feature_id": "F-001", "name": "巡检计划管理", "importance": "core", "document_sections": []}]
        )
        pp, mp, _, _ = self._write(plan, manual="# 手册\n\n软件名称：巡检系统\n\n版本号：V1.0\n\n没有功能描述。\n")
        report = cmc.run(pp, mp)
        self.assertTrue(any("巡检计划管理" in e for e in report["errors"]))

    def test_manifest_outside_plan_blocked(self):
        plan = self._plan([])
        plan["code_evidence"] = [
            {"evidence_id": "C-001", "path": "a.java", "selected": True, "sha256": "x"}
        ]
        manifest = {"files": [{"path": "b.java"}]}
        pp, _, _, mpf = self._write(plan, manifest=manifest)
        report = cmc.run(pp, manifest_path=mpf)
        self.assertTrue(any("不在已确认计划" in e for e in report["errors"]))


class FinalArtifactTest(unittest.TestCase):
    def test_docx_recheck(self):
        import docx as python_docx

        with tempfile.TemporaryDirectory() as td:
            d = Path(td)
            doc = python_docx.Document()
            doc.add_paragraph("软件名称：巡检系统")
            doc.add_paragraph("版本号：V1.0")
            doc.add_paragraph("1.1 架构")
            artifact = d / "final.docx"
            doc.save(artifact)
            report = fac.run(artifact, software_name="巡检系统", version="V1.0")
            self.assertEqual(report["status"], "pass", report["errors"])

    def test_docx_missing_name_blocked(self):
        import docx as python_docx

        with tempfile.TemporaryDirectory() as td:
            d = Path(td)
            doc = python_docx.Document()
            doc.add_paragraph("没有名称的文档")
            artifact = d / "final.docx"
            doc.save(artifact)
            report = fac.run(artifact, software_name="巡检系统", version="V1.0")
            self.assertTrue(any("软件名称" in e for e in report["errors"]))


class ProfileAndSwitchTest(unittest.TestCase):
    def test_v2_profile_recorded(self):
        sys.path.insert(0, str(SCRIPTS))
        import confirm_stage as cs

        fixture = Path(__file__).resolve().parent / "fixtures" / "boilerplate_case"
        handler = fixture / "backend" / "ImmediatePushAndAssignHandler.java"
        import hashlib

        sha = hashlib.sha256(handler.read_bytes()).hexdigest()
        with tempfile.TemporaryDirectory() as td:
            d = Path(td)
            (d / "草稿").mkdir()
            (d / "草稿" / "材料证据计划.json").write_text(
                json.dumps(
                    {
                        "schema_version": 3,
                        "software_scope": {"name": "t", "version": "v1"},
                        "input_roots": [{"root_id": "primary", "path": str(fixture)}],
                        "code_evidence": [
                            {
                                "evidence_id": "C-001",
                                "root_id": "primary",
                                "path": "backend/ImmediatePushAndAssignHandler.java",
                                "line_range": None,
                                "sha256": sha,
                                "source_kind": "first_party",
                                "grade": "A",
                                "author_declaration": {
                                    "found_author_tags": [],
                                    "categories": [],
                                    "resolution": "keep",
                                    "resolution_basis": "",
                                },
                                "selected": True,
                            }
                        ],
                        "features": [
                            {
                                "feature_id": "F-001",
                                "name": "即时推送",
                                "importance": "core",
                                "code_evidence": ["C-001"],
                                "verification": "needs_review",
                            }
                        ],
                        "fact_assertions": [],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            cs.confirm_material_plan(d, "test")
            gates = json.loads((d / "门禁状态.json").read_text(encoding="utf-8"))
            self.assertEqual(gates["material-plan"]["workflow_profile"], "v2")

    def test_switch_off_skips_cross_material(self):
        with tempfile.TemporaryDirectory() as td:
            d = Path(td)
            (d / "草稿").mkdir()
            (d / "草稿" / "材料证据计划.json").write_text(
                json.dumps({"schema_version": 3, "software_scope": {}, "features": [], "code_evidence": [], "input_roots": []}, ensure_ascii=False),
                encoding="utf-8",
            )
            (d / "门禁状态.json").write_text(
                json.dumps(
                    {
                        "material-plan": {"confirmed": True, "workflow_profile": "v2"},
                        "switches": {"cross-material": "off"},
                        "manual": {"confirmed": True},
                        "content-quality": {"confirmed": True},
                        "code-selection": {"confirmed": True},
                        "markdown": {"confirmed": True},
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            report = src.run(d)
            self.assertNotIn("cross-material", report["checks"], "开关 off 时不应执行跨材料检查")
            self.assertEqual(report["checks"]["workflow_profile"]["value"], "v2")


if __name__ == "__main__":
    unittest.main()
